from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent.parent
RACK_DIR = ROOT / "knowledge-rack"
INBOX_DIR = RACK_DIR / "inbox"
CATALOG_DIR = RACK_DIR / "catalog"
TEXT_DIR = RACK_DIR / "text"
RAW_DIR = RACK_DIR / "library" / "raw"
PUBLIC_CATALOG_DIR = ROOT / "public" / "knowledge-rack" / "catalog"
CHUNKS_PATH = CATALOG_DIR / "chunks.jsonl"
CHUNKS_JSON_PATH = CATALOG_DIR / "chunks.json"
MANIFEST_PATH = CATALOG_DIR / "manifest.json"
SUMMARY_PATH = CATALOG_DIR / "summary.md"
MAX_CHARS = 1400
PPTX_NAMESPACE = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}


@dataclass
class ExtractedDocument:
    doc_id: str
    title: str
    file_name: str
    source_type: str
    source_path: str
    raw_copy_path: str
    text_path: str
    sha256: str
    pages: int | None
    words: int
    chars: int
    chunk_count: int
    modified_at: str
    extracted_at: str


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[^a-z0-9]+", "-", value).strip("-")
    return value or "document"


def hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(path: Path) -> tuple[str, int | None]:
    page_count: int | None = None
    parts: list[str] = []
    try:
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
    except Exception:
        parts = []

    text = normalize_text("\n\n".join(parts))
    if len(text) >= 200:
        return text, page_count

    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        return text, page_count

    result = subprocess.run(
        [pdftotext, "-layout", str(path), "-"],
        capture_output=True,
        text=True,
        check=False,
    )
    fallback_text = normalize_text(result.stdout)
    return (fallback_text or text), page_count


def extract_docx_text(path: Path) -> tuple[str, int | None]:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return normalize_text("\n\n".join(paragraphs)), None


def extract_pptx_text(path: Path) -> tuple[str, int | None]:
    slides: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            slide_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            )
            for slide_name in slide_names:
                xml_bytes = archive.read(slide_name)
                root = ElementTree.fromstring(xml_bytes)
                runs = [
                    node.text.strip()
                    for node in root.findall(".//a:t", PPTX_NAMESPACE)
                    if node.text and node.text.strip()
                ]
                slide_text = normalize_text("\n".join(runs))
                if slide_text:
                    slides.append(slide_text)
            return normalize_text("\n\n".join(slides)), len(slide_names)
    except Exception:
        return "", None


def iter_source_files() -> Iterable[Path]:
    for path in sorted(INBOX_DIR.iterdir()):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() not in {".pdf", ".docx", ".pptx", ".txt", ".md"}:
            continue
        yield path


def extract_text(path: Path) -> tuple[str, int | None, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, pages = extract_pdf_text(path)
        return text, pages, "pdf"
    if suffix == ".docx":
        text, pages = extract_docx_text(path)
        return text, pages, "docx"
    if suffix == ".pptx":
        text, pages = extract_pptx_text(path)
        return text, pages, "pptx"
    return normalize_text(path.read_text(encoding="utf-8", errors="ignore")), None, "text"


def split_paragraphs(text: str) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in text.split("\n\n")]
    return [paragraph for paragraph in paragraphs if paragraph]


def chunk_text(doc_id: str, text: str) -> list[dict[str, object]]:
    paragraphs = split_paragraphs(text)
    chunks: list[dict[str, object]] = []
    current: list[str] = []
    current_len = 0

    for paragraph in paragraphs:
        extra = len(paragraph) + (2 if current else 0)
        if current and current_len + extra > MAX_CHARS:
            chunk_text_value = "\n\n".join(current).strip()
            chunks.append(
                {
                    "chunk_id": f"{doc_id}-{len(chunks) + 1:03d}",
                    "doc_id": doc_id,
                    "index": len(chunks),
                    "text": chunk_text_value,
                    "chars": len(chunk_text_value),
                    "words": len(chunk_text_value.split()),
                }
            )
            current = [paragraph]
            current_len = len(paragraph)
        else:
            current.append(paragraph)
            current_len += extra

    if current:
        chunk_text_value = "\n\n".join(current).strip()
        chunks.append(
            {
                "chunk_id": f"{doc_id}-{len(chunks) + 1:03d}",
                "doc_id": doc_id,
                "index": len(chunks),
                "text": chunk_text_value,
                "chars": len(chunk_text_value),
                "words": len(chunk_text_value.split()),
            }
        )

    return chunks


def write_summary(documents: list[ExtractedDocument]) -> None:
    lines = [
        "# Knowledge Rack Summary",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        "",
        f"Documents indexed: {len(documents)}",
        "",
        "| Title | Type | Pages | Words | Chunks |",
        "| --- | --- | ---: | ---: | ---: |",
    ]
    for document in documents:
        pages = document.pages if document.pages is not None else "-"
        lines.append(
            f"| {document.title} | {document.source_type} | {pages} | {document.words} | {document.chunk_count} |"
        )

    SUMMARY_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    CATALOG_DIR.mkdir(parents=True, exist_ok=True)
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_CATALOG_DIR.mkdir(parents=True, exist_ok=True)

    chunks: list[dict[str, object]] = []
    documents: list[ExtractedDocument] = []
    extracted_at = datetime.now(timezone.utc).isoformat()

    for source_path in iter_source_files():
        text, pages, source_type = extract_text(source_path)
        if not text:
            continue

        doc_id = slugify(source_path.stem)
        sha256 = hash_file(source_path)
        text_path = TEXT_DIR / f"{doc_id}.txt"
        raw_copy_path = RAW_DIR / source_path.name
        shutil.copy2(source_path, raw_copy_path)
        text_path.write_text(text + "\n", encoding="utf-8")

        doc_chunks = chunk_text(doc_id, text)
        for chunk in doc_chunks:
            chunk["title"] = source_path.stem
            chunk["source_path"] = str(source_path.relative_to(ROOT))
        chunks.extend(doc_chunks)

        stat = source_path.stat()
        documents.append(
            ExtractedDocument(
                doc_id=doc_id,
                title=source_path.stem,
                file_name=source_path.name,
                source_type=source_type,
                source_path=str(source_path.relative_to(ROOT)),
                raw_copy_path=str(raw_copy_path.relative_to(ROOT)),
                text_path=str(text_path.relative_to(ROOT)),
                sha256=sha256,
                pages=pages,
                words=len(text.split()),
                chars=len(text),
                chunk_count=len(doc_chunks),
                modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
                extracted_at=extracted_at,
            )
        )

    CHUNKS_PATH.write_text(
        "\n".join(json.dumps(chunk, ensure_ascii=False) for chunk in chunks) + ("\n" if chunks else ""),
        encoding="utf-8",
    )
    CHUNKS_JSON_PATH.write_text(json.dumps(chunks, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    MANIFEST_PATH.write_text(
        json.dumps([document.__dict__ for document in documents], ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    write_summary(documents)

    shutil.copyfile(MANIFEST_PATH, PUBLIC_CATALOG_DIR / "manifest.json")
    shutil.copyfile(CHUNKS_JSON_PATH, PUBLIC_CATALOG_DIR / "chunks.json")
    shutil.copyfile(SUMMARY_PATH, PUBLIC_CATALOG_DIR / "summary.md")

    print(f"Indexed {len(documents)} documents")
    print(f"Wrote manifest: {MANIFEST_PATH}")
    print(f"Wrote chunks: {CHUNKS_PATH}")
    print(f"Wrote summary: {SUMMARY_PATH}")


if __name__ == "__main__":
    main()
